"""
Экран для просмотра отчета по остаткам материалов на складе.
Отображает таблицу с текущими остатками всех материалов.
"""

import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
from config import COLORS, FONTS, SIZES
from utils import ScrollableFrame
from models import get_stock_balance


class ReportsScreen(ScrollableFrame):
    """
    Экран для отображения отчета по остаткам материалов.
    """
    
    def __init__(self, parent, app):
        super().__init__(parent, bg=COLORS['bg_main'])
        self.app = app
        self.content = self.scrollable_frame
        
        self.create_widgets()
        self.refresh_report()
    
    def create_widgets(self):
        """Создает виджеты экрана."""
        # Заголовок
        title_label = tk.Label(
            self.content,
            text="Отчет по остаткам материалов",
            font=FONTS['heading'],
            bg=COLORS['bg_main'],
            fg=COLORS['text_primary'],
            pady=20
        )
        title_label.pack()
        
        # Кнопка "Назад"
        btn_back = tk.Button(
            self.content,
            text="← Назад в меню",
            font=FONTS['body'],
            command=lambda: self.app.show_frame('MainMenu'),
            bg=COLORS['bg_button_secondary'],
            fg=COLORS['text_primary'],
            activebackground=COLORS['bg_button_hover'],
            activeforeground=COLORS['text_primary'],
            cursor="hand2",
            relief=tk.FLAT,
            padx=15,
            pady=5
        )
        btn_back.pack(pady=10)
        
        # Фрейм для кнопок управления
        control_frame = tk.Frame(self.content, bg=COLORS['bg_main'], pady=10)
        control_frame.pack()
        
        btn_refresh = tk.Button(
            control_frame,
            text="Обновить отчет",
            font=FONTS['body'],
            command=self.refresh_report,
            bg=COLORS['accent_blue'],
            fg=COLORS['text_primary'],
            activebackground='#2980b9',
            activeforeground=COLORS['text_primary'],
            cursor="hand2",
            relief=tk.FLAT,
            padx=15,
            pady=5
        )
        btn_refresh.pack(side=tk.LEFT, padx=5)
        
        btn_word = tk.Button(
            control_frame,
            text="Скачать Word",
            font=FONTS['body'],
            command=self.export_to_word,
            bg=COLORS['bg_button_success'],
            fg=COLORS['text_primary'],
            activebackground='#229954',
            activeforeground=COLORS['text_primary'],
            cursor="hand2",
            relief=tk.FLAT,
            padx=15,
            pady=5
        )
        btn_word.pack(side=tk.LEFT, padx=5)
        
        # Фрейм для таблицы
        table_frame = tk.Frame(self, bg=COLORS['bg_frame'])
        table_frame.pack(fill=tk.BOTH, expand=True, padx=SIZES['padding'], pady=SIZES['padding'])
        
        # Создаем Treeview для отображения данных
        columns = ("id", "name", "unit", "balance")
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=20)
        
        # Настраиваем колонки
        self.tree.heading("id", text="ID")
        self.tree.heading("name", text="Название материала")
        self.tree.heading("unit", text="Единица измерения")
        self.tree.heading("balance", text="Остаток")
        
        self.tree.column("id", width=80, anchor=tk.CENTER)
        self.tree.column("name", width=400, anchor=tk.W)
        self.tree.column("unit", width=200, anchor=tk.CENTER)
        self.tree.column("balance", width=200, anchor=tk.E)
        
        # Скроллбары
        v_scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # Размещаем виджеты
        self.tree.grid(row=0, column=0, sticky=tk.NSEW)
        v_scrollbar.grid(row=0, column=1, sticky=tk.NS)
        h_scrollbar.grid(row=1, column=0, sticky=tk.EW)
        
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # Метка с информацией
        self.info_label = tk.Label(
            self.content,
            text="",
            font=FONTS['body'],
            bg=COLORS['bg_main'],
            fg=COLORS['text_secondary'],
            pady=10
        )
        self.info_label.pack()
    
    def refresh_report(self):
        """Обновляет данные в отчете."""
        # Очищаем существующие записи
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Загружаем данные из БД
        balances = get_stock_balance()
        
        if not balances:
            self.info_label.config(text="Нет данных для отображения")
            return
        
        # Заполняем таблицу
        total_materials = 0
        total_balance = 0.0
        
        for balance in balances:
            material_id = balance['material_id']
            name = balance['name']
            unit = balance['unit']
            balance_value = balance['balance']
            
            # Показываем только материалы с остатком > 0
            if balance_value > 0:
                self.tree.insert(
                    "",
                    tk.END,
                    values=(
                        material_id,
                        name,
                        unit,
                        f"{balance_value:.3f}"
                    )
                )
                total_materials += 1
                total_balance += balance_value
        
        # Обновляем информацию
        if total_materials > 0:
            self.info_label.config(
                text=f"Всего материалов с остатком: {total_materials} | "
                     f"Общий остаток: {total_balance:.3f}"
            )
        else:
            self.info_label.config(text="Нет материалов с остатком на складе")
    
    def export_to_word(self):
        """Экспортирует отчет в Word документ."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            messagebox.showerror("Ошибка", "Установите python-docx: pip install python-docx")
            return
        
        try:
            # Получаем данные
            balances = get_stock_balance()
            
            # Создаем документ
            doc = Document()
            doc.add_heading('Отчет по остаткам материалов', 0)
            doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Создаем таблицу
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Заголовки
            header_cells = table.rows[0].cells
            header_cells[0].text = 'ID'
            header_cells[1].text = 'Название материала'
            header_cells[2].text = 'Единица измерения'
            header_cells[3].text = 'Остаток'
            
            # Данные
            total_balance = 0.0
            for balance in balances:
                if balance['balance'] > 0:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(balance['material_id'])
                    row_cells[1].text = balance['name']
                    row_cells[2].text = balance['unit']
                    row_cells[3].text = f"{balance['balance']:.3f}"
                    total_balance += balance['balance']
            
            # Итоговая строка
            if total_balance > 0:
                doc.add_paragraph('')
                doc.add_paragraph(f'Общий остаток: {total_balance:.3f}')
            
            # Сохраняем
            filename = f"отчет_остатки_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            messagebox.showinfo("Готово", f"Отчет сохранен: {filename}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать Word документ: {e}")
    
    def refresh(self):
        """Обновляет данные экрана при показе."""
        self.refresh_report()
